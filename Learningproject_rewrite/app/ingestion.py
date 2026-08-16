"""文档入库：解析文件 → 分块 → 写入向量库。

支持 txt / md / markdown / pdf / docx。
"""

import io
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.config import DEFAULT_CHUNK_OVERLAP, DEFAULT_CHUNK_SIZE
from app.models import IngestionResult
from app.store import delete_source_documents

TEXT_SUFFIXES = {".txt", ".md", ".markdown"}
# 按中英文标点切分，尽量避免切断语义
SPLITTER_SEPARATORS = ["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]


def parse_file(filename: str, content: bytes) -> list[Document]:
    """按扩展名解析单个文件，返回带 source 元数据的文档列表。"""
    suffix = Path(filename).suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return [
            Document(
                page_content=_decode_text(content),
                metadata={"source": filename},
            )
        ]
    if suffix == ".pdf":
        return _parse_pdf(filename, content)
    if suffix == ".docx":
        return _parse_docx(filename, content)
    raise ValueError(f"不支持的文件类型：{suffix}")


def _decode_text(raw: bytes) -> str:
    """按常见编码依次尝试解码文本文件。"""
    for encoding in ("utf-8", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _parse_pdf(filename: str, content: bytes) -> list[Document]:
    """解析 PDF：每页生成一个文档，并记录页码。"""
    reader = PdfReader(io.BytesIO(content))
    documents = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": filename, "page": page_number},
                )
            )
    return documents


def _parse_docx(filename: str, content: bytes) -> list[Document]:
    """解析 docx：按段落拼接正文。"""
    document = DocxDocument(io.BytesIO(content))
    text = "\n".join(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )
    return [Document(page_content=text, metadata={"source": filename})]


def split_documents(
    documents: list[Document],
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[Document]:
    """把文档切成适合检索的块（按中英文标点切分）。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=SPLITTER_SEPARATORS,
    )
    return splitter.split_documents(documents)


def ingest_files(
    vector_store,
    files: list[tuple[str, bytes]],
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> IngestionResult:
    """批量入库：解析、分块后写入向量库。

    同名文件会先删除旧分片再写入（即覆盖更新）；
    单个文件失败不会中断其余文件。
    """
    added_chunks = 0
    failed: list[tuple[str, str]] = []
    for filename, content in files:
        try:
            documents = parse_file(filename, content)
            if not documents:
                failed.append((filename, "未提取到文本内容"))
                continue
            chunks = split_documents(
                documents, chunk_size=chunk_size, chunk_overlap=chunk_overlap
            )
            delete_source_documents(vector_store, filename)
            if chunks:
                vector_store.add_documents(chunks)
                added_chunks += len(chunks)
        except Exception as exc:  # noqa: BLE001
            failed.append((filename, str(exc)))
    return IngestionResult(added_chunks=added_chunks, failed=failed)
